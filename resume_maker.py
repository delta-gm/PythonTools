from docx import Document  ### For python / MSWord interaction. pip install docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH  # Flags as error but isn't one


class ResumeMaker:
    
    def __init__(self, input_filename, output_filename):
        
        self.input_filename = input_filename
        self.output_filename = output_filename
        
        self.cascade = 0.3  # Horizontal spacing at start of each line (inches); makes diagonal effect
        self.spacing = 1  # Vertical spacing after lines (pt)
        self.prespace = 8  # Vertical spacing before buffer lines
        self.F0 = 'Garamond'
        self.F1 = 'Cambria'
        self.F2 = 'Times New Roman'
        self.F3 = 'Times New Roman'
        self.F4 = 'Times New Roman'
        self.F5 = 'Times New Roman'
        self.lvl0 = 20  # Text sizes
        self.lvl1 = 12
        self.lvl2 = 11
        self.lvl3 = 10
        self.lvl4 = 10
        self.lvl5 = 10
        self.grey = 100  # Grayscale color from 0 (self.black) to I think 255 (white)
        self.dark = int(self.grey / 2)
        self.black = 0
        
        self.doc = Document()  # See docx documentation. Doc is now a docx Document object and every proceeding method is from docx.
        self.sections = self.doc.sections
        self.section = self.sections[0]
        self.top_margin = 0.5  # (inches)
        self.left_margin = 0.5  # (inches)
        self.section.left_margin = Inches(self.left_margin)
        self.section.right_margin = Inches(self.left_margin)
        self.section.top_margin = Inches(self.top_margin)
        self.section.bottom_margin = Inches(self.top_margin)

    def run(self):

        ### Standard way to convert a text file into list of lines
        with open(self.input_filename) as file:
            pull_txt = file.read().splitlines()

        tags = []
        possible_tags = ['name',
                         'address',
                         'phone',
                         'email',
                         'website',
                         'summary',
                         'h1', 'h2', 'h3', 'h4', 'h5', 'h6',
                         'buffered h1', 'buffered h2',  # Heading with additional space above to denote a section
                         'end']

        
        # The relevant lines are after the tags, so create a list of numbers indicating where the tags are in the file
        for index, item in enumerate(pull_txt):
            if item in possible_tags:
                tags.append(index)

        # For each tag (instruction), format the next line of text accordingly. The list 'tags' is not the tags
        # themselves, but their line in the file.
        for index, tag in enumerate(tags):
            t = pull_txt[tag]
            if t == 'end':
                break
            if t in ['name', 'address', 'phone', 'email', 'website', 'buffer']:  # Read only the next line
                txt = pull_txt[tag + 1]
                p = self.doc.add_paragraph()
                p_f = p.paragraph_format
                p_f.space_after = Pt(self.spacing)
                if t == 'name':
                    fmat = [self.F0, self.lvl0, self.black, True, 1]
                if t in ['address', 'phone', 'email']:
                    fmat = [self.F4, self.lvl3, self.black, False, 0]
                if t == 'buffer':
                    fmat = [self.F4, self.lvl5, self.black, False, 0]
                    txt = ''
                thisrun = p.add_run(txt)
                thisrun.font.name = fmat[0]
                thisrun.font.bold = fmat[3]
                thisrun.font.small_caps = fmat[4]
                thisrun.font.size = Pt(fmat[1])
                c = fmat[2]
                thisrun.font.color.rgb = RGBColor(c, c, c)
            elif t in ['summary', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'buffered h1', 'buffered h2']:  # Read all the lines until the next tag
                txts = pull_txt[tag + 1:tags[index+1]]
                for item in txts:
                    p = self.doc.add_paragraph()
                    p_f = p.paragraph_format
                    if t == 'summary':
                        fmat = [self.F4, self.lvl3, self.black, WD_ALIGN_PARAGRAPH.CENTER, 0, False, 0, 0]
                    if t == 'h1':
                        fmat = [self.F5, self.lvl1, self.black, WD_ALIGN_PARAGRAPH.LEFT, 0, True, 2, 0]
                    if t == 'buffered h1':
                        fmat = [self.F5, self.lvl1, self.black, WD_ALIGN_PARAGRAPH.LEFT, 0, True, 2, self.prespace]
                    if t == 'h2':
                        fmat = [self.F2, self.lvl2, self.dark, WD_ALIGN_PARAGRAPH.LEFT, 1, True, 1, 0]
                    if t == 'buffered h2':
                        fmat = [self.F2, self.lvl2, self.dark, WD_ALIGN_PARAGRAPH.LEFT, 1, True, 1, self.prespace]
                    if t == 'h3':
                        fmat = [self.F2, self.lvl3, self.grey, WD_ALIGN_PARAGRAPH.LEFT, 2, True, 0, 0]
                    if t == 'h4':
                        fmat = [self.F3, self.lvl4, self.black, WD_ALIGN_PARAGRAPH.LEFT, 3, False, 0, 0]
                    if t == 'h5':
                        fmat = [self.F3, self.lvl5, self.black, WD_ALIGN_PARAGRAPH.LEFT, 4, False, 0, 0]
                        p.style = 'List Bullet'
                    thisrun = p.add_run(item)
                    if fmat[6] == 0:
                        thisrun.font.all_caps = False
                        thisrun.font.small_caps = False
                    elif fmat[6] == 1:
                        thisrun.font.small_caps = True
                    else:
                        thisrun.font.all_caps = True
                    thisrun.font.name = fmat[0]
                    thisrun.font.size = Pt(fmat[1])
                    c = fmat[2]
                    p_f.alignment = fmat[3]
                    p_f.space_after = Pt(self.spacing)
                    p_f.space_before = Pt(fmat[7])
                    p_f.left_indent = Inches(fmat[4] * self.cascade)
                    thisrun.font.color.rgb = RGBColor(c, c, c)
                    thisrun.font.bold = fmat[5]

        self.doc.save(self.output_filename)


RM = ResumeMaker('emi_resume_text.txt', 'Emis_Resume.docx')
RM.run()
